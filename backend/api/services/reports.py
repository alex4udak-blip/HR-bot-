import io
import os
import re
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Font configuration - fonts are in api/fonts/, not api/services/fonts/
FONTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'fonts')
_fonts_registered = False
_use_cyrillic_fonts = False

def _register_fonts():
    """Register DejaVu fonts for Cyrillic support, fallback to Helvetica."""
    global _fonts_registered, _use_cyrillic_fonts
    if _fonts_registered:
        return _use_cyrillic_fonts

    _fonts_registered = True

    # Try to register DejaVu fonts
    font_path = os.path.join(FONTS_DIR, 'DejaVuSans.ttf')
    font_bold_path = os.path.join(FONTS_DIR, 'DejaVuSans-Bold.ttf')

    if os.path.exists(font_path) and os.path.exists(font_bold_path):
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', font_bold_path))
            # Register font family for bold/italic mapping
            addMapping('DejaVuSans', 0, 0, 'DejaVuSans')  # normal
            addMapping('DejaVuSans', 1, 0, 'DejaVuSans-Bold')  # bold
            addMapping('DejaVuSans', 0, 1, 'DejaVuSans')  # italic (use regular)
            addMapping('DejaVuSans', 1, 1, 'DejaVuSans-Bold')  # bold italic
            _use_cyrillic_fonts = True
            print("DejaVu fonts registered successfully")
            return True
        except Exception as e:
            print(f"Warning: Could not register DejaVu fonts: {e}")
    else:
        print(f"Warning: Font files not found at {FONTS_DIR}")

    _use_cyrillic_fonts = False
    return False


def generate_pdf_report(title: str, content: str, chat_title: str) -> bytes:
    """Generate PDF report from markdown content."""
    use_cyrillic = _register_fonts()

    # Choose fonts based on availability
    font_regular = 'DejaVuSans' if use_cyrillic else 'Helvetica'
    font_bold = 'DejaVuSans-Bold' if use_cyrillic else 'Helvetica-Bold'

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm, leftMargin=20*mm, rightMargin=20*mm)

    styles = getSampleStyleSheet()

    # Custom styles with appropriate fonts
    styles.add(ParagraphStyle(
        name='CustomTitle',
        fontName=font_bold,
        fontSize=18,
        spaceAfter=12,
        textColor=colors.HexColor('#1a1a2e'),
    ))
    styles.add(ParagraphStyle(
        name='CustomHeading',
        fontName=font_bold,
        fontSize=14,
        spaceAfter=8,
        spaceBefore=16,
        textColor=colors.HexColor('#2d2d44'),
    ))
    styles.add(ParagraphStyle(
        name='CustomSubheading',
        fontName=font_bold,
        fontSize=12,
        spaceAfter=6,
        spaceBefore=12,
        textColor=colors.HexColor('#3d3d54'),
    ))
    styles.add(ParagraphStyle(
        name='CustomBody',
        fontName=font_regular,
        fontSize=10,
        spaceAfter=6,
        leading=14,
        textColor=colors.HexColor('#333344'),
    ))
    styles.add(ParagraphStyle(
        name='CustomBullet',
        fontName=font_regular,
        fontSize=10,
        spaceAfter=4,
        leftIndent=20,
        bulletIndent=10,
        leading=14,
        textColor=colors.HexColor('#333344'),
    ))
    styles.add(ParagraphStyle(
        name='CustomQuote',
        fontName=font_regular,
        fontSize=9,
        spaceAfter=6,
        leftIndent=15,
        leading=12,
        textColor=colors.HexColor('#666677'),
        borderLeftWidth=2,
        borderLeftColor=colors.HexColor('#ccccdd'),
    ))

    story = []

    # Title
    story.append(Paragraph(escape_html(f"Аналитический отчёт: {chat_title}"), styles['CustomTitle']))
    story.append(Paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['CustomBody']))
    story.append(Spacer(1, 20))

    # Parse markdown and add content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 8))
        elif line.startswith('### '):
            clean = escape_html(line[4:])
            story.append(Paragraph(clean, styles['CustomSubheading']))
        elif line.startswith('## '):
            clean = escape_html(line[3:])
            story.append(Paragraph(clean, styles['CustomHeading']))
        elif line.startswith('# '):
            clean = escape_html(line[2:])
            story.append(Paragraph(clean, styles['CustomTitle']))
        elif line.startswith('> '):
            # Quote
            clean = format_markdown(line[2:])
            story.append(Paragraph(f"« {clean} »", styles['CustomQuote']))
        elif line.startswith('**') and line.endswith('**'):
            # Bold line
            clean = escape_html(line[2:-2])
            story.append(Paragraph(f"<b>{clean}</b>", styles['CustomBody']))
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point
            text = line[2:]
            clean = format_markdown(text)
            story.append(Paragraph(f"• {clean}", styles['CustomBullet']))
        elif line.startswith('* '):
            # Bullet point with asterisk
            clean = format_markdown(line[2:])
            story.append(Paragraph(f"• {clean}", styles['CustomBullet']))
        elif re.match(r'^\d+\.\s', line):
            # Numbered list
            clean = format_markdown(line)
            story.append(Paragraph(clean, styles['CustomBullet']))
        elif line.startswith('---') or line.startswith('___'):
            # Horizontal rule - add spacer
            story.append(Spacer(1, 15))
        else:
            # Regular text with markdown formatting
            clean = format_markdown(line)
            if clean:  # Only add non-empty paragraphs
                story.append(Paragraph(clean, styles['CustomBody']))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def format_markdown(text: str) -> str:
    """Convert markdown formatting to ReportLab XML tags."""
    # Escape HTML first
    text = escape_html(text)
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Remove any remaining problematic characters
    text = text.replace('```', '')
    return text


def generate_docx_report(title: str, content: str, chat_title: str) -> bytes:
    """Generate DOCX report from markdown content."""
    doc = Document()

    # Title
    title_para = doc.add_heading(f"Аналитический отчёт: {chat_title}", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Parse and add content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        elif line.startswith('---') or line.startswith('___'):
            # Horizontal rule - add empty paragraph
            doc.add_paragraph()
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('> '):
            # Quote - add as indented italic text
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run(f"« {line[2:]} »")
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point
            text = line[2:]
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, text)
        elif line.startswith('* '):
            # Bullet point with asterisk
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, line[2:])
        elif re.match(r'^\d+\.\s', line):
            # Numbered list
            para = doc.add_paragraph(style='List Number')
            # Remove the number prefix and add text
            text = re.sub(r'^\d+\.\s', '', line)
            _add_formatted_text(para, text)
        else:
            # Regular paragraph with formatting
            para = doc.add_paragraph()
            _add_formatted_text(para, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_formatted_text(para, text: str):
    """Add text to paragraph with bold/italic formatting from markdown."""
    # Pattern to match **bold** and *italic*
    pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic text
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            para.add_run(part)
