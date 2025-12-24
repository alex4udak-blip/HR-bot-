"""
Comprehensive unit tests for miscellaneous services:
- entity_ai.py: AI assistant for entity analysis
- reports.py: PDF and DOCX report generation
- password_policy.py: Password validation
"""
import pytest
import pytest_asyncio
from datetime import datetime
from unittest.mock import AsyncMock, MagicMock, patch, PropertyMock, mock_open
from typing import List, AsyncGenerator
import io
import os

from api.services.entity_ai import (
    EntityAIService,
    ENTITY_QUICK_ACTIONS,
    entity_ai_service
)
from api.services.reports import (
    generate_pdf_report,
    generate_docx_report,
    escape_html,
    format_markdown,
    _add_formatted_text,
    _register_fonts
)
from api.services.password_policy import (
    validate_password,
    COMMON_PASSWORDS
)
from api.models.database import (
    Entity,
    Chat,
    Message,
    CallRecording,
    EntityType,
    EntityStatus,
    ChatType,
    CallStatus,
    CallSource
)


# ============================================================================
# ENTITY AI SERVICE TESTS
# ============================================================================

class TestEntityAIService:
    """Tests for EntityAIService."""

    @pytest.fixture
    def mock_settings(self):
        """Mock settings with API key."""
        with patch('api.services.entity_ai.settings') as mock_settings:
            mock_settings.anthropic_api_key = "test-api-key-12345"
            yield mock_settings

    @pytest.fixture
    def mock_settings_no_key(self):
        """Mock settings without API key."""
        with patch('api.services.entity_ai.settings') as mock_settings:
            mock_settings.anthropic_api_key = None
            yield mock_settings

    @pytest.fixture
    def sample_entity(self) -> Entity:
        """Create sample entity for testing."""
        entity = Entity(
            id=1,
            org_id=1,
            department_id=1,
            created_by=1,
            name="John Doe",
            email="john@example.com",
            phone="+1234567890",
            company="ACME Corp",
            position="Senior Developer",
            type=EntityType.candidate,
            status=EntityStatus.active,
            tags=["python", "senior"],
            created_at=datetime(2025, 1, 1, 10, 0, 0)
        )
        return entity

    @pytest.fixture
    def sample_chat(self) -> Chat:
        """Create sample chat for testing."""
        chat = Chat(
            id=1,
            org_id=1,
            owner_id=1,
            telegram_chat_id=123456789,
            title="HR Interview",
            custom_name="Interview with John",
            chat_type=ChatType.hr,
            is_active=True,
            created_at=datetime(2025, 1, 1, 10, 0, 0)
        )

        # Add messages
        chat.messages = [
            Message(
                id=1,
                chat_id=1,
                telegram_message_id=1001,
                telegram_user_id=101,
                username="john_doe",
                first_name="John",
                last_name="Doe",
                content="Hello, I'm interested in the position.",
                content_type="text",
                timestamp=datetime(2025, 1, 1, 10, 5, 0)
            ),
            Message(
                id=2,
                chat_id=1,
                telegram_message_id=1002,
                telegram_user_id=102,
                username="hr_manager",
                first_name="HR",
                last_name="Manager",
                content="Great! Tell me about your experience.",
                content_type="text",
                timestamp=datetime(2025, 1, 1, 10, 10, 0)
            ),
            Message(
                id=3,
                chat_id=1,
                telegram_message_id=1003,
                telegram_user_id=101,
                username="john_doe",
                first_name="John",
                last_name="Doe",
                content="I have 5 years of Python development experience.",
                content_type="text",
                timestamp=datetime(2025, 1, 1, 10, 15, 0)
            ),
        ]

        return chat

    @pytest.fixture
    def sample_call(self) -> CallRecording:
        """Create sample call recording for testing."""
        call = CallRecording(
            id=1,
            org_id=1,
            owner_id=1,
            title="Technical Interview",
            source_type=CallSource.upload,
            status=CallStatus.done,
            duration_seconds=1800,
            transcript="Interviewer: Tell me about your Python experience.\nCandidate: I've been working with Python for 5 years...",
            summary="Discussed candidate's Python experience and past projects.",
            key_points=[
                "5 years of Python experience",
                "Worked on microservices architecture",
                "Experience with FastAPI and Django"
            ],
            created_at=datetime(2025, 1, 2, 14, 0, 0)
        )
        return call

    def test_entity_ai_service_init(self):
        """Test EntityAIService initialization."""
        service = EntityAIService()
        assert service._client is None
        assert service.model == "claude-sonnet-4-20250514"

    def test_client_property_with_api_key(self, mock_settings):
        """Test client property when API key is configured."""
        service = EntityAIService()
        service._client = None  # Reset client

        with patch('api.services.entity_ai.AsyncAnthropic') as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.return_value = mock_client

            client = service.client

            assert client is not None
            mock_anthropic.assert_called_once_with(api_key="test-api-key-12345")

    def test_client_property_without_api_key(self, mock_settings_no_key):
        """Test client property raises error when API key is missing."""
        service = EntityAIService()
        service._client = None  # Reset client

        with pytest.raises(ValueError, match="ANTHROPIC_API_KEY не настроен"):
            _ = service.client

    def test_build_entity_context_basic(self, sample_entity):
        """Test building entity context with basic info."""
        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [], [])

        assert "John Doe" in context
        assert "candidate" in context
        assert "active" in context
        assert "ACME Corp" in context
        assert "Senior Developer" in context
        assert "john@example.com" in context
        assert "+1234567890" in context
        assert "python, senior" in context
        assert "К этому контакту пока не привязаны чаты или звонки" in context

    def test_build_entity_context_with_chats(self, sample_entity, sample_chat):
        """Test building entity context with chat messages."""
        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [sample_chat], [])

        assert "ПЕРЕПИСКИ:" in context
        assert "Interview with John" in context
        assert "Hello, I'm interested in the position." in context
        assert "Great! Tell me about your experience." in context
        assert "I have 5 years of Python development experience." in context
        assert "John Doe" in context
        assert "HR Manager" in context

    def test_build_entity_context_with_calls(self, sample_entity, sample_call):
        """Test building entity context with call recordings."""
        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [], [sample_call])

        assert "ЗВОНКИ:" in context
        assert "Technical Interview" in context
        assert "30м 0с" in context  # Duration
        assert "Discussed candidate's Python experience" in context
        assert "5 years of Python experience" in context
        assert "Worked on microservices architecture" in context
        assert "Tell me about your Python experience" in context

    def test_build_entity_context_with_long_messages(self, sample_entity, sample_chat):
        """Test building entity context truncates long messages."""
        # Add a very long message
        long_message = Message(
            id=4,
            chat_id=1,
            telegram_message_id=1004,
            telegram_user_id=101,
            username="john_doe",
            first_name="John",
            last_name="Doe",
            content="A" * 1000,  # Very long content
            content_type="text",
            timestamp=datetime(2025, 1, 1, 10, 20, 0)
        )
        sample_chat.messages.append(long_message)

        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [sample_chat], [])

        # Message should be truncated to 500 chars
        assert "A" * 500 in context
        assert "A" * 501 not in context

    def test_build_entity_context_with_long_transcript(self, sample_entity, sample_call):
        """Test building entity context truncates long transcripts."""
        # Create a call with very long transcript
        sample_call.transcript = "B" * 6000

        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [], [sample_call])

        # Transcript should be truncated to 5000 chars
        assert "B" * 5000 in context
        assert "(транскрипт обрезан)" in context

    def test_build_entity_context_message_without_name(self, sample_entity, sample_chat):
        """Test building entity context with message without name."""
        # Add message without first/last name or username
        message = Message(
            id=5,
            chat_id=1,
            telegram_message_id=1005,
            telegram_user_id=105,
            username=None,
            first_name=None,
            last_name=None,
            content="Test message",
            content_type="text",
            timestamp=datetime(2025, 1, 1, 10, 25, 0)
        )
        sample_chat.messages.append(message)

        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [sample_chat], [])

        assert "Unknown" in context
        assert "Test message" in context

    def test_build_entity_context_media_message(self, sample_entity, sample_chat):
        """Test building entity context with media message without content."""
        # Add media message without text content
        message = Message(
            id=6,
            chat_id=1,
            telegram_message_id=1006,
            telegram_user_id=101,
            username="john_doe",
            first_name="John",
            last_name="Doe",
            content=None,  # No text content
            content_type="photo",
            timestamp=datetime(2025, 1, 1, 10, 30, 0)
        )
        sample_chat.messages.append(message)

        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [sample_chat], [])

        assert "[медиа]" in context

    def test_build_system_prompt(self, sample_entity):
        """Test building system prompt."""
        service = EntityAIService()
        context = service._build_entity_context(sample_entity, [], [])
        system_prompt = service._build_system_prompt(context)

        assert "AI-ассистент для HR-аналитики" in system_prompt
        assert context in system_prompt
        assert "Отвечай на русском языке" in system_prompt
        assert "ТОЛЬКО на фактах" in system_prompt
        assert "конкретные цитаты" in system_prompt

    @pytest.mark.asyncio
    async def test_chat_stream_success(self, mock_settings, sample_entity, sample_chat, sample_call):
        """Test successful streaming chat."""
        service = EntityAIService()

        # Mock the Anthropic client
        mock_stream = AsyncMock()
        mock_stream.text_stream = AsyncIterator(["Hello", " ", "world", "!"])
        mock_stream.__aenter__ = AsyncMock(return_value=mock_stream)
        mock_stream.__aexit__ = AsyncMock(return_value=None)

        mock_messages = MagicMock()
        mock_messages.stream = MagicMock(return_value=mock_stream)

        mock_client = MagicMock()
        mock_client.messages = mock_messages

        service._client = mock_client

        # Test streaming
        result = []
        async for chunk in service.chat_stream(
            "Tell me about this person",
            sample_entity,
            [sample_chat],
            [sample_call],
            []
        ):
            result.append(chunk)

        assert result == ["Hello", " ", "world", "!"]
        mock_messages.stream.assert_called_once()

    @pytest.mark.asyncio
    async def test_chat_stream_with_history(self, mock_settings, sample_entity):
        """Test streaming chat with conversation history."""
        service = EntityAIService()

        # Mock the Anthropic client
        mock_stream = AsyncMock()
        mock_stream.text_stream = AsyncIterator(["Response"])
        mock_stream.__aenter__ = AsyncMock(return_value=mock_stream)
        mock_stream.__aexit__ = AsyncMock(return_value=None)

        mock_messages = MagicMock()
        mock_messages.stream = MagicMock(return_value=mock_stream)

        mock_client = MagicMock()
        mock_client.messages = mock_messages

        service._client = mock_client

        # Test with history
        history = [
            {"role": "user", "content": "Previous question"},
            {"role": "assistant", "content": "Previous answer"}
        ]

        result = []
        async for chunk in service.chat_stream(
            "Follow-up question",
            sample_entity,
            [],
            [],
            history
        ):
            result.append(chunk)

        # Verify the messages parameter included history
        call_kwargs = mock_messages.stream.call_args[1]
        messages = call_kwargs['messages']
        assert len(messages) == 3  # 2 history + 1 new
        assert messages[0]["content"] == "Previous question"
        assert messages[1]["content"] == "Previous answer"
        assert messages[2]["content"] == "Follow-up question"

    @pytest.mark.asyncio
    async def test_chat_stream_error(self, mock_settings, sample_entity):
        """Test streaming chat error handling."""
        service = EntityAIService()

        # Mock the Anthropic client to raise an error
        mock_client = MagicMock()
        mock_client.messages.stream.side_effect = Exception("API Error")
        service._client = mock_client

        with pytest.raises(Exception, match="API Error"):
            async for _ in service.chat_stream("Test", sample_entity, [], [], []):
                pass

    @pytest.mark.asyncio
    async def test_quick_action_success(self, mock_settings, sample_entity):
        """Test successful quick action execution."""
        service = EntityAIService()

        # Mock chat_stream
        async def mock_stream(*args, **kwargs):
            yield "Analysis"
            yield " result"

        with patch.object(service, 'chat_stream', side_effect=mock_stream):
            result = []
            async for chunk in service.quick_action("full_analysis", sample_entity, [], []):
                result.append(chunk)

            assert result == ["Analysis", " result"]

    @pytest.mark.asyncio
    async def test_quick_action_unknown(self, sample_entity):
        """Test quick action with unknown action."""
        service = EntityAIService()

        result = []
        async for chunk in service.quick_action("unknown_action", sample_entity, [], []):
            result.append(chunk)

        assert len(result) == 1
        assert "Неизвестное действие: unknown_action" in result[0]

    @pytest.mark.asyncio
    async def test_quick_action_all_types(self, mock_settings, sample_entity):
        """Test all quick action types are available."""
        service = EntityAIService()

        # Mock chat_stream
        async def mock_stream(*args, **kwargs):
            yield "Result"

        with patch.object(service, 'chat_stream', side_effect=mock_stream):
            for action in ENTITY_QUICK_ACTIONS.keys():
                result = []
                async for chunk in service.quick_action(action, sample_entity, [], []):
                    result.append(chunk)

                assert len(result) > 0

    def test_get_available_actions(self):
        """Test getting available quick actions."""
        service = EntityAIService()
        actions = service.get_available_actions()

        assert len(actions) == 6
        action_ids = [a["id"] for a in actions]
        assert "full_analysis" in action_ids
        assert "red_flags" in action_ids
        assert "comparison" in action_ids
        assert "prediction" in action_ids
        assert "summary" in action_ids
        assert "questions" in action_ids

        # Check structure
        for action in actions:
            assert "id" in action
            assert "label" in action
            assert "icon" in action

    def test_entity_ai_service_singleton(self):
        """Test entity_ai_service singleton instance."""
        assert entity_ai_service is not None
        assert isinstance(entity_ai_service, EntityAIService)


# ============================================================================
# REPORTS SERVICE TESTS
# ============================================================================

class TestReportsService:
    """Tests for report generation functions."""

    @pytest.fixture
    def sample_markdown(self) -> str:
        """Sample markdown content for testing."""
        return """# Main Heading

## Section 1

This is a paragraph with **bold text** and *italic text*.

### Subsection

- Bullet point 1
- Bullet point 2 with **bold**
* Alternative bullet
• Unicode bullet

1. Numbered item 1
2. Numbered item 2

> This is a quote

**Standalone bold line**

---

Regular text after horizontal rule.
"""

    def test_escape_html(self):
        """Test HTML escaping."""
        assert escape_html("Hello World") == "Hello World"
        assert escape_html("<script>alert('xss')</script>") == "&lt;script&gt;alert('xss')&lt;/script&gt;"
        assert escape_html("A & B") == "A &amp; B"
        assert escape_html("3 < 5 > 2") == "3 &lt; 5 &gt; 2"

    def test_format_markdown_bold(self):
        """Test markdown bold formatting."""
        result = format_markdown("This is **bold** text")
        assert result == "This is <b>bold</b> text"

    def test_format_markdown_italic(self):
        """Test markdown italic formatting."""
        result = format_markdown("This is *italic* text")
        assert result == "This is <i>italic</i> text"

    def test_format_markdown_combined(self):
        """Test combined markdown formatting."""
        result = format_markdown("This is **bold** and *italic* text")
        assert result == "This is <b>bold</b> and <i>italic</i> text"

    def test_format_markdown_escapes_html(self):
        """Test that markdown formatting also escapes HTML."""
        result = format_markdown("<script>**test**</script>")
        assert "<script>" not in result
        assert "&lt;script&gt;" in result
        assert "<b>test</b>" in result

    def test_format_markdown_removes_code_blocks(self):
        """Test that code block markers are removed."""
        result = format_markdown("```python\ncode\n```")
        assert "```" not in result
        assert "python" in result
        assert "code" in result

    def test_register_fonts_success(self):
        """Test font registration success."""
        with patch('api.services.reports.os.path.exists', return_value=True):
            with patch('api.services.reports.pdfmetrics.registerFont'):
                with patch('api.services.reports.addMapping'):
                    with patch('api.services.reports.TTFont') as mock_ttfont:
                        # Reset the global flag
                        import api.services.reports as reports_module
                        reports_module._fonts_registered = False

                        result = _register_fonts()

                        assert result is True
                        assert reports_module._use_cyrillic_fonts is True

    def test_register_fonts_not_found(self):
        """Test font registration when fonts are not found."""
        with patch('api.services.reports.os.path.exists', return_value=False):
            # Reset the global flag
            import api.services.reports as reports_module
            reports_module._fonts_registered = False

            result = _register_fonts()

            assert result is False
            assert reports_module._use_cyrillic_fonts is False

    def test_register_fonts_already_registered(self):
        """Test that fonts are only registered once."""
        import api.services.reports as reports_module
        reports_module._fonts_registered = True
        reports_module._use_cyrillic_fonts = True

        with patch('api.services.reports.pdfmetrics.registerFont') as mock_register:
            result = _register_fonts()

            # Should not call registerFont again
            mock_register.assert_not_called()
            assert result is True

    def test_register_fonts_error(self):
        """Test font registration error handling."""
        with patch('api.services.reports.os.path.exists', return_value=True):
            with patch('api.services.reports.pdfmetrics.registerFont', side_effect=Exception("Font error")):
                with patch('api.services.reports.TTFont'):
                    # Reset the global flag
                    import api.services.reports as reports_module
                    reports_module._fonts_registered = False

                    result = _register_fonts()

                    assert result is False

    def test_generate_pdf_report_basic(self, sample_markdown):
        """Test PDF report generation with basic content."""
        with patch('api.services.reports._register_fonts', return_value=True):
            pdf_bytes = generate_pdf_report(
                title="Test Report",
                content=sample_markdown,
                chat_title="Test Chat"
            )

            assert isinstance(pdf_bytes, bytes)
            assert len(pdf_bytes) > 0
            # PDF files start with %PDF
            assert pdf_bytes.startswith(b'%PDF')

    def test_generate_pdf_report_cyrillic(self):
        """Test PDF report with Cyrillic characters."""
        content = "# Тестовый отчёт\n\nЭто **русский** текст."

        with patch('api.services.reports._register_fonts', return_value=True):
            pdf_bytes = generate_pdf_report(
                title="Отчёт",
                content=content,
                chat_title="Русский чат"
            )

            assert isinstance(pdf_bytes, bytes)
            assert len(pdf_bytes) > 0

    def test_generate_pdf_report_without_cyrillic_fonts(self, sample_markdown):
        """Test PDF report generation without Cyrillic fonts."""
        with patch('api.services.reports._register_fonts', return_value=False):
            pdf_bytes = generate_pdf_report(
                title="Test Report",
                content=sample_markdown,
                chat_title="Test Chat"
            )

            assert isinstance(pdf_bytes, bytes)
            assert len(pdf_bytes) > 0

    def test_generate_pdf_report_empty_content(self):
        """Test PDF report with empty content."""
        with patch('api.services.reports._register_fonts', return_value=True):
            pdf_bytes = generate_pdf_report(
                title="Empty Report",
                content="",
                chat_title="Empty Chat"
            )

            assert isinstance(pdf_bytes, bytes)
            assert len(pdf_bytes) > 0

    def test_generate_docx_report_basic(self, sample_markdown):
        """Test DOCX report generation with basic content."""
        docx_bytes = generate_docx_report(
            title="Test Report",
            content=sample_markdown,
            chat_title="Test Chat"
        )

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        # DOCX files are ZIP archives, start with PK
        assert docx_bytes.startswith(b'PK')

    def test_generate_docx_report_cyrillic(self):
        """Test DOCX report with Cyrillic characters."""
        content = "# Тестовый отчёт\n\nЭто **русский** текст с *курсивом*."

        docx_bytes = generate_docx_report(
            title="Отчёт",
            content=content,
            chat_title="Русский чат"
        )

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    def test_generate_docx_report_empty_content(self):
        """Test DOCX report with empty content."""
        docx_bytes = generate_docx_report(
            title="Empty Report",
            content="",
            chat_title="Empty Chat"
        )

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    def test_add_formatted_text_bold(self):
        """Test adding bold formatted text to DOCX."""
        from docx import Document
        doc = Document()
        para = doc.add_paragraph()

        _add_formatted_text(para, "This is **bold** text")

        runs = para.runs
        assert len(runs) == 3
        assert runs[0].text == "This is "
        assert runs[0].bold is not True
        assert runs[1].text == "bold"
        assert runs[1].bold is True
        assert runs[2].text == " text"
        assert runs[2].bold is not True

    def test_add_formatted_text_italic(self):
        """Test adding italic formatted text to DOCX."""
        from docx import Document
        doc = Document()
        para = doc.add_paragraph()

        _add_formatted_text(para, "This is *italic* text")

        runs = para.runs
        assert len(runs) == 3
        assert runs[0].text == "This is "
        assert runs[1].text == "italic"
        assert runs[1].italic is True
        assert runs[2].text == " text"

    def test_add_formatted_text_combined(self):
        """Test adding combined formatted text to DOCX."""
        from docx import Document
        doc = Document()
        para = doc.add_paragraph()

        _add_formatted_text(para, "This is **bold** and *italic* text")

        runs = para.runs
        assert len(runs) == 5
        assert runs[1].bold is True
        assert runs[3].italic is True

    def test_add_formatted_text_no_formatting(self):
        """Test adding plain text to DOCX."""
        from docx import Document
        doc = Document()
        para = doc.add_paragraph()

        _add_formatted_text(para, "Plain text")

        runs = para.runs
        assert len(runs) == 1
        assert runs[0].text == "Plain text"


# ============================================================================
# PASSWORD POLICY TESTS
# ============================================================================

class TestPasswordPolicy:
    """Tests for password policy validation."""

    def test_validate_password_valid_basic(self):
        """Test valid password with basic requirements."""
        is_valid, error = validate_password("Test1234")
        assert is_valid is True
        assert error == ""

    def test_validate_password_valid_complex(self):
        """Test valid password with complex characters."""
        is_valid, error = validate_password("MyP@ssw0rd123!")
        assert is_valid is True
        assert error == ""

    def test_validate_password_too_short(self):
        """Test password that is too short."""
        is_valid, error = validate_password("Test12")
        assert is_valid is False
        assert "at least 8 characters" in error

    def test_validate_password_no_number(self):
        """Test password without a number."""
        is_valid, error = validate_password("TestPassword")
        assert is_valid is False
        assert "at least one number" in error

    def test_validate_password_no_letter(self):
        """Test password without a letter."""
        is_valid, error = validate_password("12345678")
        assert is_valid is False
        assert "at least one letter" in error

    def test_validate_password_common_password(self):
        """Test common passwords are rejected."""
        common_passwords = ["password", "123456", "password123", "qwerty", "admin"]

        for pwd in common_passwords:
            is_valid, error = validate_password(pwd)
            assert is_valid is False
            assert "too common" in error.lower()

    def test_validate_password_common_case_insensitive(self):
        """Test common password check is case-insensitive."""
        is_valid, error = validate_password("PASSWORD")
        assert is_valid is False
        assert "too common" in error.lower()

    def test_validate_password_matches_email(self):
        """Test password cannot be same as email."""
        is_valid, error = validate_password("test@example.com", email="test@example.com")
        assert is_valid is False
        assert "cannot be same as email" in error

    def test_validate_password_matches_email_case_insensitive(self):
        """Test password email check is case-insensitive."""
        is_valid, error = validate_password("TEST@EXAMPLE.COM", email="test@example.com")
        assert is_valid is False
        assert "cannot be same as email" in error

    def test_validate_password_matches_email_username(self):
        """Test password cannot be same as email username."""
        is_valid, error = validate_password("testuser1", email="testuser@example.com")
        assert is_valid is False
        assert "cannot be same as email username" in error

    def test_validate_password_matches_email_username_case_insensitive(self):
        """Test password email username check is case-insensitive."""
        is_valid, error = validate_password("TESTUSER1", email="testuser@example.com")
        assert is_valid is False
        assert "cannot be same as email username" in error

    def test_validate_password_no_email_provided(self):
        """Test password validation without email."""
        is_valid, error = validate_password("GoodPass123")
        assert is_valid is True
        assert error == ""

    def test_validate_password_empty(self):
        """Test empty password."""
        is_valid, error = validate_password("")
        assert is_valid is False
        assert "at least 8 characters" in error

    def test_validate_password_special_characters(self):
        """Test password with special characters is valid."""
        is_valid, error = validate_password("P@ssw0rd!#$")
        assert is_valid is True
        assert error == ""

    def test_validate_password_unicode(self):
        """Test password with Unicode characters."""
        is_valid, error = validate_password("Пароль123")
        assert is_valid is True
        assert error == ""

    def test_validate_password_boundary_length(self):
        """Test password at boundary length."""
        # Exactly 8 characters
        is_valid, error = validate_password("Pass1234")
        assert is_valid is True
        assert error == ""

        # 7 characters
        is_valid, error = validate_password("Pass123")
        assert is_valid is False
        assert "at least 8 characters" in error

    def test_validate_password_multiple_numbers(self):
        """Test password with multiple numbers."""
        is_valid, error = validate_password("Test12345678")
        assert is_valid is True
        assert error == ""

    def test_validate_password_multiple_letters(self):
        """Test password with multiple letters."""
        is_valid, error = validate_password("TestPassword1")
        assert is_valid is True
        assert error == ""

    def test_common_passwords_set(self):
        """Test COMMON_PASSWORDS set contains expected passwords."""
        assert "password" in COMMON_PASSWORDS
        assert "123456" in COMMON_PASSWORDS
        assert "admin" in COMMON_PASSWORDS
        assert "qwerty" in COMMON_PASSWORDS
        assert len(COMMON_PASSWORDS) > 30  # Should have a good list

    def test_validate_password_similar_to_email_but_valid(self):
        """Test password similar to email but not matching is valid."""
        is_valid, error = validate_password("TestUser123", email="testuser@example.com")
        # Should be valid because it has numbers added
        assert is_valid is True
        assert error == ""


# ============================================================================
# HELPER CLASS FOR ASYNC ITERATION
# ============================================================================

class AsyncIterator:
    """Helper class for async iteration in tests."""

    def __init__(self, items):
        self.items = items
        self.index = 0

    def __aiter__(self):
        return self

    async def __anext__(self):
        if self.index >= len(self.items):
            raise StopAsyncIteration
        item = self.items[self.index]
        self.index += 1
        return item
